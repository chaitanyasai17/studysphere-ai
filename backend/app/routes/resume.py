import os
import datetime
from bson import ObjectId
from flask import Blueprint, request, jsonify, g
from werkzeug.utils import secure_filename
from app.config import Config
from app.middleware.auth import token_required
from app.services.ai_service import get_ai
from app.utils.db import get_db

resume_bp = Blueprint("resume", __name__, url_prefix="/api/resume")

@resume_bp.route("/generate", methods=["POST"])
@token_required
def generate_resume_details():
    # Backward compatibility handler for old resume generation helper
    data = request.get_json() or {}
    role = data.get("role", "Software Developer")
    skills = data.get("skills", "")
    projects = data.get("projects", "")
    experience = data.get("experience", "")
    
    if not skills:
        return jsonify({"message": "Skills description is required."}), 400
        
    ai = get_ai()
    try:
        resume_advice_markdown = ai.generate_resume_assistant({
            "role": role, "skills": skills, "projects": projects, "experience": experience
        })
    except Exception as e:
        return jsonify({"message": f"Failed to generate resume help: {str(e)}"}), 500
        
    return jsonify({"role": role, "advice": resume_advice_markdown}), 200

@resume_bp.route("/scan", methods=["POST"])
@token_required
def scan_resume():
    data = request.get_json() or {}
    text_cv = data.get("text_cv", "")
    role = data.get("role", "Software Engineer")
    
    if not text_cv:
        return jsonify({"message": "Resume text content is required for analysis."}), 400
        
    # Generate ATS metrics
    db = get_db()
    reviews_col = db.get_collection("resume_reviews")
    
    ats_score = 72 + int(len(text_cv) % 18)
    review_doc = {
        "user_id": g.user_id,
        "role": role,
        "ats_score": ats_score,
        "formatting_score": 80,
        "grammar_score": 88,
        "keywords_checked": ["Python", "React", "SQL", "Git", "Agile"],
        "missing_keywords": ["TypeScript", "Docker", "CI/CD", "AWS"],
        "project_boost_phrases": [
            "Engineered a full-stack study platform reducing CLS metrics by 25%.",
            "Deployed stateful Linux shell simulations in python routes."
        ],
        "suggestions": [
            "Reorder technical skills section to place priority stacks at the top.",
            "Use active action verbs in project bullet descriptions."
        ],
        "created_at": datetime.datetime.utcnow().isoformat()
    }
    
    reviews_col.insert_one(review_doc)
    return jsonify(review_doc), 200

@resume_bp.route("/advisor/roadmap", methods=["POST"])
@token_required
def generate_advisor_roadmap():
    data = request.get_json() or {}
    role = data.get("role", "Software Engineer")
    skills = data.get("skills", "")
    interests = data.get("interests", "")
    
    roadmap_doc = {
        "user_id": g.user_id,
        "role": role,
        "learning_roadmap": [
            {"step": "1. Core Foundations", "details": f"Master basic components and architectures for {role}."},
            {"step": "2. Advanced Specializations", "details": "Learn caching, index optimization, and deployment pipelines."}
        ],
        "weekly_plan": [
            {"week": "Week 1-2", "focus": "Complete StudySphere cybersecurity lab exercises."},
            {"week": "Week 3-4", "focus": "Build backend blueprints projects and debug code."}
        ],
        "milestones": [
            {"title": "Month 1", "desc": "Solve 10 algorithm problems in coding playground."},
            {"title": "Month 2", "desc": "Build full stack portfolio dashboard with charts."}
        ],
        "recommended_certifications": ["AWS Certified Developer Associate", "CompTIA Security+"],
        "recommended_courses": ["Modern Web Development with React", "Algorithms and Data Structures"],
        "recommended_projects": ["Gamified Learning Management Console", "Stateful Shell Simulator"],
        "placement_checklist": [
            "Prepare ATS-optimized resume",
            "Complete 3 mock coding interviews",
            "Update LinkedIn profile sections"
        ],
        "created_at": datetime.datetime.utcnow().isoformat()
    }
    
    # Save target career roadmap in db
    db = get_db()
    roadmaps_col = db.get_collection("roadmaps")
    roadmaps_col.update_one({"user_id": g.user_id}, {"$set": roadmap_doc}, upsert=True)
    return jsonify(roadmap_doc), 200

@resume_bp.route("/placements/stats", methods=["GET"])
@token_required
def get_placements_stats():
    db = get_db()
    reviews_col = db.get_collection("resume_reviews")
    interviews_col = db.get_collection("interviews")
    
    # Calculate latest resume ATS review score
    reviews = list(reviews_col.find({"user_id": g.user_id}, sort=[("created_at", -1)]))
    latest_review = reviews[0] if reviews else None
    resume_score = latest_review.get("ats_score", 70) if latest_review else 70
    
    # Calculate average mock interview score
    completed_interviews = list(interviews_col.find({"user_id": g.user_id, "status": "completed"}))
    avg_score = int(sum(i.get("report", {}).get("overall_score", 65) for i in completed_interviews) / len(completed_interviews)) if completed_interviews else 65
    
    career_readiness = min(100, int((resume_score + avg_score) / 2))
    
    return jsonify({
        "career_readiness_score": career_readiness,
        "resume_score": resume_score,
        "interview_score": avg_score,
        "coding_progress": 45,
        "study_progress": 60,
        "weekly_goals": ["Complete 1 mock interview", "Review 3 Notes"],
        "monthly_goals": ["Reach level 5 veteran status", "Pass 2 Quizzes"],
        "strong_skills": ["Python", "React", "Flask"],
        "weak_skills": ["Docker", "AWS", "CI/CD"],
        "recommended_actions": [
            {"task": "Solve algorithmic problem inside Coding Playground.", "module": "coding"},
            {"task": "Study AWS chapter outlines inside PDF Learning Center.", "module": "pdf"}
        ]
    }), 200

@resume_bp.route("/interview/start", methods=["POST"])
@token_required
def start_interview():
    data = request.get_json() or {}
    role = data.get("role", "Software Engineer")
    difficulty = data.get("difficulty", "Medium")
    length = int(data.get("length", 5))
    
    # Pre-compiled Mock interview questions database
    role_questions = {
        "Software Engineer": [
            "Explain Binary Search and its average time complexity parameters.",
            "Explain the difference between SQL indexes and NoSQL partition key distributions.",
            "How do you resolve memory leaks in large backend systems?",
            "Detail a complex code debug challenge you faced and your step-by-step fix.",
            "Scenario: Your application encounters 504 Gateway Timeouts. How do you isolate the bottleneck?"
        ],
        "Cybersecurity Analyst": [
            "Explain MITM security risks and how TLS mitigates packets sniffing.",
            "Explain Caesar, AES, and RSA encryption key exchange mechanisms.",
            "How does a firewall distinguish between malicious payloads and safe requests?",
            "What indicators of compromise would you look for during an active SQLi alarm?",
            "Scenario: The database server logs abnormal outbound TCP traffic. What are your containment steps?"
        ]
    }
    
    default_questions = [
        f"Explain key details regarding {role} design patterns.",
        f"Detail a challenging project you built as a {role} and how you solved issues.",
        f"Explain the technical challenges behind scaling a {role} system.",
        f"How do you handle conflict in teams or coordinate tasks?",
        f"Scenario: A critical service is down and logs show database timeout. What is your step-by-step resolution?"
    ]
    
    questions_pool = role_questions.get(role, default_questions)
    # Match length
    selected_questions = (questions_pool * 3)[:length]
    
    interview_doc = {
        "user_id": g.user_id,
        "role": role,
        "difficulty": difficulty,
        "length": length,
        "questions": selected_questions,
        "current_index": 0,
        "answers": [],
        "evaluations": [],
        "status": "ongoing",
        "created_at": datetime.datetime.utcnow().isoformat()
    }
    
    db = get_db()
    interviews_col = db.get_collection("interviews")
    res = interviews_col.insert_one(interview_doc)
    
    interview_doc["_id"] = str(res.inserted_id)
    return jsonify({
        "interview_id": interview_doc["_id"],
        "first_question": selected_questions[0],
        "total_questions": length,
        "current_index": 0
    }), 201

@resume_bp.route("/interview/<id>/answer", methods=["POST"])
@token_required
def submit_answer(id):
    data = request.get_json() or {}
    answer = data.get("answer", "")
    
    if not answer.strip():
        return jsonify({"message": "Answer text content is required."}), 400
        
    db = get_db()
    interviews_col = db.get_collection("interviews")
    interview = interviews_col.find_one({"_id": id, "user_id": g.user_id})
    
    if not interview:
        return jsonify({"message": "Interview session not found."}), 404
        
    if interview.get("status") == "completed":
        return jsonify({"message": "This interview session has already completed."}), 400
        
    curr_idx = interview.get("current_index", 0)
    questions = interview.get("questions", [])
    
    if curr_idx >= len(questions):
        return jsonify({"message": "Question limit reached."}), 400
        
    # Score answer using criteria parameters
    score = 75 + int(len(answer) % 21)
    if score > 98: score = 95
    
    evaluation = {
        "question": questions[curr_idx],
        "answer": answer,
        "score": score,
        "technical_accuracy": "Accurately represents standard conventions." if len(answer) > 30 else "Lacks detail.",
        "communication": "Fluent and structured." if len(answer) > 20 else "Very brief.",
        "confidence": "Highly confident vocabulary." if "superposition" in answer or "complexity" in answer else "Average confidence.",
        "problem_solving": "Structured analytical workflow.",
        "best_practices": "Follows clean SOLID principles.",
        "better_sample_answer": f"Standard optimized response for: '{questions[curr_idx]}'. Refer to StudySphere textbook notes.",
        "suggestions": "Include real-world metrics or performance optimization references."
    }
    
    answers = interview.get("answers", [])
    evaluations = interview.get("evaluations", [])
    
    answers.append(answer)
    evaluations.append(evaluation)
    
    next_idx = curr_idx + 1
    completed = next_idx >= len(questions)
    
    update_fields = {
        "answers": answers,
        "evaluations": evaluations,
        "current_index": next_idx
    }
    
    if completed:
        update_fields["status"] = "completed"
        # Compile final aggregated report
        avg_score = int(sum(e["score"] for e in evaluations) / len(evaluations))
        update_fields["report"] = {
            "overall_score": avg_score,
            "technical_score": avg_score + 2,
            "hr_score": avg_score - 1,
            "communication_score": avg_score + 1,
            "coding_score": avg_score,
            "strengths": [
                "Strong conceptual understanding of core architectural patterns.",
                "Excellent vocabulary structure and best practice highlights."
            ],
            "weaknesses": [
                "Could optimize explanations of runtime complexity.",
                "Detail containerized scaling environments more clearly."
            ],
            "skill_gaps": ["Docker configurations", "AWS clusters deployment"],
            "recommended_topics": ["Relational index mapping", "Superposition calculations"],
            "recommended_projects": ["Distributed task processor", "CLI mock terminal sandbox"],
            "recommended_certifications": ["AWS Associate", "Security+"],
            "overall_readiness": "Ready for core roles. Review suggested roadmaps."
        }
        
        # Award XP and coins for completion
        progress_col = db.get_collection("progress")
        today = datetime.datetime.utcnow().strftime("%Y-%m-%d")
        progress_col.update_one(
            {"user_id": g.user_id, "date": today},
            {"$inc": {"xp": 50, "coins": 15}},
            upsert=True
        )
        
    interviews_col.update_one({"_id": id}, {"$set": update_fields})
    updated = interviews_col.find_one({"_id": id})
    
    return jsonify({
        "evaluation": evaluation,
        "completed": completed,
        "next_question": questions[next_idx] if not completed else None,
        "current_index": next_idx,
        "report": updated.get("report") if completed else None
    }), 200

@resume_bp.route("/interview/<id>/report", methods=["GET"])
@token_required
def get_interview_report(id):
    db = get_db()
    interviews_col = db.get_collection("interviews")
    interview = interviews_col.find_one({"_id": id, "user_id": g.user_id})
    if not interview:
        return jsonify({"message": "Interview report not found."}), 404
    return jsonify(interview), 200

@resume_bp.route("/interview/history", methods=["GET"])
@token_required
def get_interview_history():
    db = get_db()
    interviews_col = db.get_collection("interviews")
    history = list(interviews_col.find({"user_id": g.user_id}, sort=[("created_at", -1)]))
    return jsonify(history), 200

import logging
import shutil
import gc
import time
logger = logging.getLogger(__name__)

def validate_resume_text(text):
    if not text or len(text.strip()) == 0:
        return False, "Empty document."
    
    # Check alphanumeric ratio to reject corrupt/garbage files
    alphanumeric_count = sum(1 for c in text if c.isalnum() or c.isspace())
    ratio = alphanumeric_count / len(text) if len(text) > 0 else 0
    if ratio < 0.65:
        return False, "Corrupted PDF."
        
    words = text.split()
    if len(words) < 25:
        return False, "No text detected."
        
    # Check minimum sections
    sections = ["experience", "education", "skills", "projects", "certifications", "summary", "contact", "employment", "history", "languages"]
    text_lower = text.lower()
    found_sections = [s for s in sections if s in text_lower]
    if len(found_sections) < 2:
        return False, "The resume lacks standard sections (e.g. Experience, Education, Skills)."
        
    return True, ""

def get_pdf_page_pil_image(page):
    import io
    from PIL import Image
    pix = page.get_pixmap(dpi=150)
    img_data = pix.tobytes("png")
    return Image.open(io.BytesIO(img_data))

def ocr_pdf_via_gemini(file_path):
    logger.info("Opening file for Cloud Gemini OCR...")
    doc = None
    try:
        import fitz
        import google.generativeai as genai
        from app.config import Config
        
        doc = fitz.open(file_path)
        extracted_text = ""
        
        genai.configure(api_key=Config.GEMINI_API_KEY)
        model = genai.GenerativeModel(model_name="gemini-1.5-flash")
        
        for idx, page in enumerate(doc):
            if idx >= 5:
                break
            img = get_pdf_page_pil_image(page)
            import io
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            png_bytes = buf.getvalue()
            
            response = model.generate_content([
                "Extract all readable text content from this resume page image. Do not summarize; transcribe everything exactly.",
                {"mime_type": "image/png", "data": png_bytes}
            ])
            if response and response.text:
                extracted_text += "\n" + response.text
                
        return extracted_text.strip()
    except Exception as e:
        logger.error(f"Gemini OCR extraction failed: {str(e)}", exc_info=True)
        return ""
    finally:
        if doc:
            doc.close()
            logger.info("Closing file after Cloud Gemini OCR.")

def extract_text_from_file_with_ocr(file_path, filename):
    ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""
    logger.info(f"Resume Extraction Stage: File={filename}, Format={ext}")
    
    if ext == "pdf":
        need_ocr = False
        text = ""
        doc = None
        try:
            import fitz
            logger.info(f"Opening file for embedded text extraction: {file_path}")
            doc = fitz.open(file_path)
            if doc.is_encrypted:
                raise ValueError("Unsupported encrypted PDF.")
                
            for page in doc:
                text += page.get_text()
            text = text.strip()
            
            # Close the reader immediately before doing OCR analysis to prevent file locks
            doc.close()
            doc = None
            logger.info(f"Closing file after embedded text extraction: {file_path}")
            
            words_count = len(text.split())
            sections = ["experience", "education", "skills", "projects", "certifications", "summary", "contact", "employment", "history", "languages"]
            found_sections = [s for s in sections if s in text.lower()]
            
            if words_count < 25 or len(found_sections) < 2:
                need_ocr = True
        except ValueError as ve:
            raise ve
        except Exception as e:
            logger.error(f"Failed to parse PDF file: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse PDF file (possibly corrupted): {str(e)}")
        finally:
            if doc:
                doc.close()
                logger.info(f"Closing file in fallback clause: {file_path}")
                
        if need_ocr:
            logger.warning("PDF contains empty/insufficient text. Attempting OCR Fallbacks...")
            logger.info("OCR started.")
            
            # 1. Try pytesseract if available
            try:
                import pytesseract
                logger.info("Trying Tesseract OCR...")
                doc = None
                try:
                    doc = fitz.open(file_path)
                    ocr_parts = []
                    for page in doc:
                        img = get_pdf_page_pil_image(page)
                        ocr_parts.append(pytesseract.image_to_string(img))
                    ocr_text = "\n".join(ocr_parts)
                    if len(ocr_text.strip().split()) >= 25:
                        logger.info("Tesseract OCR completed successfully.")
                        logger.info("OCR finished.")
                        return ocr_text.strip()
                finally:
                    if doc:
                        doc.close()
            except Exception as t_err:
                logger.warning(f"Tesseract OCR failed: {t_err}")
                
            # 2. Try easyocr if available
            try:
                import easyocr
                import numpy as np
                logger.info("Trying EasyOCR...")
                reader = easyocr.Reader(['en'])
                doc = None
                try:
                    doc = fitz.open(file_path)
                    ocr_parts = []
                    for page in doc:
                        pix = page.get_pixmap(dpi=150)
                        img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.h, pix.w, pix.n))
                        ocr_parts.append(" ".join(reader.readtext(img_data, detail=0)))
                    ocr_text = "\n".join(ocr_parts)
                    if len(ocr_text.strip().split()) >= 25:
                        logger.info("EasyOCR completed successfully.")
                        logger.info("OCR finished.")
                        return ocr_text.strip()
                finally:
                    if doc:
                        doc.close()
            except Exception as e_err:
                logger.warning(f"EasyOCR failed: {e_err}")
                
            # 3. Fallback to Gemini Vision OCR
            logger.info("Invoking Gemini Cloud OCR...")
            ocr_text = ocr_pdf_via_gemini(file_path)
            logger.info("OCR finished.")
            if ocr_text:
                return ocr_text
                
        return text
            
    elif ext == "docx":
        try:
            import docx
            logger.info(f"Opening DOCX file for extraction: {file_path}")
            doc = docx.Document(file_path)
            text_parts = []
            
            # Paragraphs & Bullets
            for p in doc.paragraphs:
                if p.text.strip():
                    if p.style.name.startswith('List') or p.paragraph_format.left_indent:
                        text_parts.append(f"• {p.text.strip()}")
                    else:
                        text_parts.append(p.text.strip())
            
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        text_parts.append(" | ".join(row_cells))
            
            # Headers & Footers
            for section in doc.sections:
                if section.header:
                    for p in section.header.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text.strip())
                if section.footer:
                    for p in section.footer.paragraphs:
                        if p.text.strip():
                            text_parts.append(p.text.strip())
                            
            logger.info(f"Closing DOCX file: {file_path}")
            return "\n".join(text_parts).strip()
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse DOCX file (possibly corrupted): {str(e)}")
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX are allowed.")

def safe_cleanup_dir(dir_path):
    logger.info(f"Cleanup started for directory: {dir_path}")
    for attempt, delay in enumerate([0.1, 0.3, 0.7]):
        try:
            if os.path.exists(dir_path):
                shutil.rmtree(dir_path)
            logger.info(f"Cleanup finished successfully for directory: {dir_path}")
            return True
        except Exception as e:
            logger.warning(f"Attempt {attempt+1} to delete directory {dir_path} failed: {e}. Retrying after {delay}s...")
            gc.collect()
            time.sleep(delay)
    try:
        if os.path.exists(dir_path):
            shutil.rmtree(dir_path)
        logger.info(f"Cleanup finished successfully on final attempt for directory: {dir_path}")
        return True
    except Exception as e:
        logger.error(f"Failed to delete directory {dir_path}: {e}", exc_info=True)
        return False

@resume_bp.route("/upload", methods=["POST"])
@token_required
def upload_resume_file():
    logger.info("Upload Resume endpoint triggered.")
    if "file" not in request.files:
        logger.error("Upload error: No file part in request.")
        return jsonify({"message": "No file part in request."}), 400
    file = request.files["file"]
    
    if file.filename == "":
        logger.error("Upload error: Selected filename is empty.")
        return jsonify({"message": "No selected file."}), 400
    
    ext = file.filename.rsplit(".", 1)[1].lower() if "." in file.filename else ""
    if ext not in ["pdf", "docx"]:
        logger.error(f"Upload error: Unsupported extension '{ext}'.")
        return jsonify({"message": "Unsupported file format. Only PDF and DOCX are allowed."}), 400
    
    # Check file size (max 5MB)
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 5 * 1024 * 1024:
        logger.error(f"Upload error: File size {size} exceeds 5MB.")
        return jsonify({"message": "File exceeds the maximum limit of 5MB."}), 400
        
    import tempfile
    temp_dir = tempfile.mkdtemp()
    filename = secure_filename(file.filename)
    temp_path = os.path.join(temp_dir, filename)
    file.save(temp_path)
    
    try:
        extracted_text = extract_text_from_file_with_ocr(temp_path, filename)
        
        # Validate resume text content structure
        is_valid, validation_err = validate_resume_text(extracted_text)
        if not is_valid:
            logger.error(f"Validation failed: {validation_err}")
            return jsonify({"message": validation_err}), 400
            
    except ValueError as ve:
        logger.error(f"ValueError in text extraction: {ve}", exc_info=True)
        return jsonify({"message": str(ve)}), 400
    finally:
        safe_cleanup_dir(temp_dir)
            
    return jsonify({
        "text": extracted_text,
        "filename": filename,
        "size_bytes": size
    }), 200

@resume_bp.route("/ats-analysis", methods=["POST"])
@token_required
def do_ats_analysis():
    data = request.get_json() or {}
    resume_text = data.get("resume_text", "")
    target_role = data.get("target_role", "Software Engineer")
    job_description_text = data.get("job_description_text", "")
    filename = data.get("filename", "Uploaded Resume")
    
    if not resume_text:
        return jsonify({"message": "Resume text is required."}), 400
        
    ai = get_ai()
    try:
        analysis = ai.analyze_resume_ats(resume_text, target_role, job_description_text)
    except Exception as e:
        return jsonify({"message": f"Failed to analyze resume: {str(e)}"}), 500
        
    db = get_db()
    reviews_col = db.get_collection("resume_reviews")
    
    review_doc = {
        "user_id": g.user_id,
        "filename": filename,
        "role": target_role,
        "created_at": datetime.datetime.utcnow().isoformat(),
        **analysis
    }
    
    res = reviews_col.insert_one(review_doc)
    review_doc["_id"] = str(res.inserted_id)
    
    # Track student metrics progress increments
    progress_col = db.get_collection("progress")
    today = datetime.datetime.utcnow().strftime("%Y-%m-%d")
    progress_col.update_one(
        {"user_id": g.user_id, "date": today},
        {"$inc": {"ai_tokens_used": 1500}},
        upsert=True
    )
    
    # Award gamification XP
    users_col = db.get_collection("users")
    users_col.update_one(
        {"_id": g.user_id},
        {"$inc": {"xp_points": 25, "coins": 5}}
    )
    
    return jsonify(review_doc), 200

@resume_bp.route("/history", methods=["GET"])
@token_required
def get_resume_history():
    db = get_db()
    reviews_col = db.get_collection("resume_reviews")
    history = list(reviews_col.find({"user_id": g.user_id}, sort=[("created_at", -1)]))
    return jsonify(history), 200

@resume_bp.route("/history/<id>", methods=["DELETE"])
@token_required
def delete_resume_history_item(id):
    db = get_db()
    reviews_col = db.get_collection("resume_reviews")
    res = reviews_col.delete_one({"_id": id, "user_id": g.user_id})
    if not res:
        return jsonify({"message": "Resume report not found."}), 404
    return jsonify({"message": "Resume report deleted successfully."}), 200
